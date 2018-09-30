from gensim.models import Word2Vec, KeyedVectors
from gensim.models.phrases import Phrases, Phraser
#from pattern3 import es
#import textract
from os import listdir
from os.path import isfile, join
import pickle
import PyPDF2
import os, sys
import docx
import re
from nltk.stem.wordnet import WordNetLemmatizer
import nltk
import math

def getText(filename):
	with open(filename, 'r') as file:
		data = file.read()
	return data
	
def getTextPDF(filename):
	pdfFileObj = open(filename, 'rb')
	pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
	num_pages = pdfReader.numPages
	count = 0
	text = ""
	#The while loop will read each page
	while count < num_pages:
		pageObj = pdfReader.getPage(count)
		count +=1
		text += pageObj.extractText()
	return text

def getTextDOCX(filename):
	doc = docx.Document(filename)
	fullText = []
	for para in doc.paragraphs:
		fullText.append(para.text)
	return ' '.join(fullText)

def preprocess(text):
	text = re.sub('[^A-Za-z0-9]', ' ', text)
	text = ' '.join(text.split())
	return text

def orgPDF(filename):
    pdfFileObj = open(filename, 'rb')
    pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
    num_pages = pdfReader.numPages
    count = 0
    text = []
    #The while loop will read each page
    while count < num_pages:
        pageObj = pdfReader.getPage(count)
        count +=1
        text += pageObj.extractText().splitlines()
    pdfFileObj.close()
    return text

def orgTXT(filename):
    file = open(filename, 'r', encoding='utf-8')
    text = file.read().splitlines()
    return text

def orgDOC(filename):
    file = open(filename, 'r', encoding='utf-8')
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return fullText


# =============================================================================
# filename = 'C:/Users/Sarthak/Desktop/Hack/a2.pdf'
# if filename.endswith('.pdf'):
# 	text = getTextPDF(filename)
# elif filename.endswith('.docx'):
# 	text = getTextDOCX(filename)
# text = preprocess(text)
# print(text)
# 
# =============================================================================
def preprocess_training_data1(dir_cvs):    
    dircvs = [join(dir_cvs, f) for f in listdir(dir_cvs) if isfile(join(dir_cvs, f))]
    alltext = []
    print('Reading resumes')
    for cv in dircvs:
        print(cv)
        if cv.endswith('.pdf'):
            try:
                text = getTextPDF(cv)
            except:
                continue
        elif cv.endswith('.docx'):
            try:
                text = getTextDOCX(cv)
            except:
                continue
        elif cv.endswith('.txt'):
            try:
                text = getText(cv)
            except:
                continue
        else:
            continue
        text = preprocess(text)
        text = text.lower()
        text = (' ').join([word for word in text.split(' ') if len(word)>2])
        alltext.append(text)    

    print('Pre-processing data')
    lmtzr = WordNetLemmatizer()
    nouns = []
    for sentence in alltext:
        temp = []
        for word,pos in nltk.pos_tag(nltk.word_tokenize(str(sentence))):
            # word = stemmer.stem(word)
            if (pos == 'NN' or pos == 'NNP' or pos == 'NNS' or pos == 'NNPS' or pos == 'VB'):
                temp.append(lmtzr.lemmatize(word))
        nouns.append(temp)
    print(nouns)
    return nouns


def train_word2vec(vector, dir_model_name):
    model = Word2Vec(vector, size=200, window=5, min_count=3, workers=4)
    model.save(dir_model_name)
#    print(model['ibm'])
    return model
    

def train_bigrams(vector, dir_model_name):
    bigrams = Phrases(vector)
    model = Word2Vec(bigrams[vector],  size=150, window=5, min_count=3, workers=4)
    model.save(dir_model_name)
#    print(model)
    return model

def get_trained_model(model):
    if os.path.exists(model):
        return Word2Vec.load(model)
    else:
        data = preprocess_training_data1('JobX/gensim/ResumeDataset_new')
        data = [x for x in data if x != []]
        new_model = train_bigrams(data, model)
        return new_model

def preprocess_word(word):
    lmtzr = WordNetLemmatizer()
    return lmtzr.lemmatize(word)

def extract_keywords(cv, new_model, keywords, top_n=20, thresh=0.97):
    if cv.endswith('.pdf'):
        # try:
        text = getTextPDF(cv)
        # except:
        #     print('Error occured while parsing the resume')
        #     pass
    elif cv.endswith('.docx') or cv.endswith('.doc'):
        # try:
        text = getTextDOCX(cv)
        # except:
        #     print('Error occured while parsing the resume')
        #     pass
    elif cv.endswith('.txt'):
        # try:
        text = getText(cv)
        # except:
            # pass
    else:
        return []
    
    text = preprocess(text)
    text = text.lower()
    text = (' ').join([word for word in text.split(' ') if len(word)>2])
#    print(text)
    skills = []
    text = text.split(' ')
    for word in text:
        flag = 0
        word1 = preprocess_word(word)
        for keyword in keywords:
            word2 = preprocess_word(keyword)
            try:
                if(new_model.wv.similarity(word1, word2) < thresh):
                    flag = 1
                    break
            except:
                flag = 1
                break
        if flag==0:
            skills.append(word)
        else:
            continue
    key_points = []
    try:
        try:
            org_text = orgPDF(cv)
        except:
            org_text = orgDOC(cv)
    except:
        org_text = orgTXT(cv)
    for skill in skills:
        key_points += [sentence + '.' for sentence in org_text if skill in sentence.lower()]
#    print(org_text)
    seen = set()
    seen_add = seen.add
    summary = [x for x in key_points if not (x in seen or seen_add(x))]
    
    s = set()
    s_add = s.add
    skills = [x for x in skills if not (x in s or s_add(x))]
    return summary, skills
#    return key_points, skills
    
    ############################# JD ####################################
    
def generateSummary(filename, jd, keywords=['strength', 'skill'], thresh=0.6, addfilePath=True):
    if addfilePath:
        filename = "JobX/uploads/"+filename
    print(filename, file=sys.stderr)
    new_model = get_trained_model('JobX/gensim/bigram_trained_model_new.bin')
    ans = extract_keywords(filename, new_model, keywords, thresh=0.7)
    skills=ans[1]
    jd = jd.split(' ')
    jd_skills = []
    for word in jd:
        flag = 0
        word1 = preprocess_word(word)
        for keyword in keywords:
            word2 = preprocess_word(keyword)
            try:
                if(new_model.wv.similarity(word1, word2) < thresh):
                    flag = 1
                    break
            except:
                flag = 1
                break
        if flag==0:
            jd_skills.append(word)
        else:
            continue
    s = set()
    s_add = s.add
    jd_skills = [x for x in jd_skills if not (x in s or s_add(x))]
    
    for candidate in skills:
        word1 = preprocess_word(candidate)
        score = 0.0
        for req in jd_skills:
            word2 = preprocess_word(req)
            try:
                score += new_model.wv.similarity(word1, word2)
            except:
                pass
    try:
        final_score = float(score)/len(skills)
        final_score = final_score/math.ceil(final_score)
    except:
        final_score = 0.0      
    return ans[0], ans[1], jd_skills, final_score
        

def generateSummaryCV(filename):
    filename = "JobX/uploads/"+filename
    new_model = get_trained_model('JobX/gensim/bigram_trained_model_new.bin')
    ans = extract_keywords(filename, new_model, [ 'strength', 'skill'], thresh=0.7)
# print(ans[0])
# print()
# print()
# print(ans[1])
# #print(get_score(jd, ans[1], new_model, [ 'strength', 'skill'] ))
# print()
    
    
    




